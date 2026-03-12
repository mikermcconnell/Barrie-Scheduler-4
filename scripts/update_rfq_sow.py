from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE_PATH = r"C:\Users\Mike McConnell\Downloads\11756 Professional Services RFQ (SOW).docx"


def insert_paragraph_after(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def main():
    doc = Document(SOURCE_PATH)
    paragraphs = list(doc.paragraphs)

    replacements = {
        9: (
            "The primary objective of this project is to identify, evaluate, and develop concept designs for a safe, efficient, and future-ready transit hub at the Barrie Campus. The project will include a campus-wide location screening process to establish a consultant-led shortlist of three candidate hub locations, developed in consultation with Georgian College staff and Barrie Transit staff. For the shortlisted locations, the project will assess transportation flow for buses, pedestrians, cyclists, and vehicles while improving accessibility, safety, and user experience. The study will culminate in three conceptual design options to support campus growth, sustainability goals, and integration with the City of Barrie's and the County of Simcoe's broader transit network."
        ),
        61: (
            "This section requires comprehensive assessment of existing conditions for pedestrians, cyclists, bus, and vehicle movements with safety analysis to support the evaluation of candidate hub locations and the detailed review of the shortlisted options."
        ),
        62: (
            "Study Area: The consultant shall consider the broader Barrie Campus and any adjacent access and interface areas necessary to identify, compare, and recommend candidate hub locations. Following shortlisting, detailed existing conditions analysis shall focus on the three shortlisted locations, associated access routes, and connections to key campus buildings, transit interfaces, and active transportation corridors."
        ),
        129: (
            "This section requires development of conceptual design drawings for the three shortlisted bus hub location options, including platforms, shelters, lighting, and supporting infrastructure, with integrated improvement recommendations from the location screening and existing movement and safety assessment. The functional design must be developed enough to confirm the feasibility of each shortlisted location and provide sufficient detail for accurate costing."
        ),
        131: (
            "Creates detailed design drawings for three shortlisted hub location options incorporating all required infrastructure elements, amenities, and systems to support hub operations."
        ),
        133: (
            "Three complete conceptual design drawing sets (one for each shortlisted location option)"
        ),
        155: (
            "Develops specific infrastructure improvements for pedestrians, cyclists, transit, and vehicle movements based on the safety assessment findings to create a safe and efficient way to access and egress each shortlisted hub location within the applicable study area."
        ),
        180: (
            "This section addresses consideration of future campus planning to ensure the preferred hub design accommodates potential development opportunities. Future proofing analysis is required only for the preferred design option identified through the consultant's evaluation of the shortlisted locations and concept options."
        ),
        195: (
            "Assesses the feasibility and design implications of accommodating articulated buses in the future for the preferred design option only, including platform extensions, circulation requirements, and any site modifications needed at the preferred shortlisted location."
        ),
        208: (
            "This section requires Class D cost estimates for all three shortlisted design options plus annual maintenance cost projections and comparative analysis."
        ),
        210: (
            "Provides detailed construction cost estimates for each shortlisted design option with appropriate contingencies and supporting documentation for budget planning."
        ),
        212: (
            "Individual cost estimates for each of the three shortlisted design options"
        ),
        233: (
            "Annual maintenance cost projections for each shortlisted design option up to 10 years post occupancy"
        ),
        249: (
            "Campus-wide transportation master planning beyond the level required to identify, screen, and compare candidate hub locations"
        ),
        254: (
            "Analysis of campus interior roadways beyond what is required to assess access to the shortlisted hub locations"
        ),
        329: (
            "ANNEX B - Existing Study Area Map (Background Reference)"
        ),
        330: (
            "ANNEX C - Existing Sponsor Concept Layouts (Background Reference)"
        ),
    }

    for index, new_text in replacements.items():
        paragraphs[index].text = new_text

    responsibilities_anchor = paragraphs[18]
    insert_paragraph_after(
        responsibilities_anchor,
        "Availability of Georgian College staff and Barrie Transit staff for candidate location review and shortlisting sessions",
        "List Paragraph",
    )

    location_heading = insert_paragraph_after(
        paragraphs[58],
        "Candidate Location Identification & Shortlisting",
        "Heading 2",
    )
    location_intro = insert_paragraph_after(
        location_heading,
        "This section requires the consultant to identify, screen, and recommend candidate mobility hub locations across Barrie Campus and any adjacent interface areas needed for transit access, rather than limiting the study to the east side concepts previously identified.",
        "Normal",
    )
    screening_heading = insert_paragraph_after(
        location_intro,
        "Location Screening & Consultation",
        "Heading 3",
    )
    screening_intro = insert_paragraph_after(
        screening_heading,
        "Establishes a consultant-led process to review the full range of viable hub locations, confirm evaluation criteria with project stakeholders, and narrow the list to three shortlisted locations for concept development.",
        "Normal",
    )
    deliverables_heading = insert_paragraph_after(
        screening_intro,
        "Deliverables",
        "Heading 3",
    )

    last = deliverables_heading
    for text in [
        "Candidate location inventory and screening matrix",
        "Evaluation criteria and weighting methodology",
        "Summary of consultation input from Georgian College staff and Barrie Transit staff",
        "Recommended shortlist of three candidate locations with rationale",
    ]:
        last = insert_paragraph_after(last, text, "List Paragraph")

    technical_heading = insert_paragraph_after(last, "Technical Requirements", "Heading 3")
    last = technical_heading
    for text in [
        "Review campus-wide location opportunities and constraints, including transit operations, pedestrian and cyclist access, campus circulation, servicing, constructability, and future development interface",
        "Minimum two working sessions with Georgian College staff and Barrie Transit staff to review the long list and confirm the shortlist",
        "Screening criteria must include safety, transit operability, accessibility, customer convenience, campus integration, implementation complexity, and future expansion potential",
        "Existing sponsor concepts are to be reviewed as background information only and shall not be treated as the only candidate locations",
    ]:
        last = insert_paragraph_after(last, text, "List Paragraph")

    acceptance_heading = insert_paragraph_after(last, "Acceptance Criteria", "Heading 3")
    last = acceptance_heading
    for text in [
        "Three shortlisted locations must be clearly supported by documented evaluation criteria and stakeholder input",
        "The shortlist must identify the advantages, constraints, and rationale for advancing each location",
        "Recommended locations must be suitable for subsequent concept design and cost estimation",
    ]:
        last = insert_paragraph_after(last, text, "List Paragraph")

    doc.save(SOURCE_PATH)


if __name__ == "__main__":
    main()
