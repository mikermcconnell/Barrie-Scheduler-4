from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE_PATH = r"C:\Users\Mike McConnell\Downloads\11756 Professional Services RFQ (SOW).docx"
OUTPUT_PATH = r"C:\Users\Mike McConnell\Downloads\11756 Professional Services RFQ (SOW) - marked changes.docx"

ADDED_PARAGRAPHS = {
    20,
    61, 62, 63, 64, 65, 66, 67, 68, 69,
    70, 71, 72, 73, 74, 75, 76, 77, 78,
}

REVISED_PARAGRAPHS = {
    10, 81, 82, 149, 151, 153, 175, 200, 215,
    228, 230, 232, 253, 269, 274, 349, 350,
}


def insert_paragraph_after(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def rewrite_with_marker(paragraph, marker, fill_color, marker_rgb):
    original_text = paragraph.text
    paragraph.text = ""

    marker_run = paragraph.add_run(marker)
    marker_run.bold = True
    if marker_rgb is not None:
        marker_run.font.color.rgb = marker_rgb
    marker_run.font.highlight_color = fill_color

    content_run = paragraph.add_run(original_text)
    content_run.font.highlight_color = fill_color


def main():
    doc = Document(SOURCE_PATH)
    paragraphs = list(doc.paragraphs)

    title_note = insert_paragraph_after(
        paragraphs[0],
        "Review copy: added paragraphs are marked [ADDED] and highlighted green. Revised existing paragraphs are marked [REVISED] and highlighted yellow.",
        "Normal",
    )
    rewrite_with_marker(
        title_note,
        "[NOTE] ",
        WD_COLOR_INDEX.GRAY_25,
        None,
    )

    from docx.shared import RGBColor

    green = RGBColor(0x00, 0x61, 0x00)
    red = RGBColor(0x9C, 0x00, 0x06)

    for index in sorted(ADDED_PARAGRAPHS):
        rewrite_with_marker(
            paragraphs[index - 1],
            "[ADDED] ",
            WD_COLOR_INDEX.BRIGHT_GREEN,
            green,
        )

    for index in sorted(REVISED_PARAGRAPHS):
        rewrite_with_marker(
            paragraphs[index - 1],
            "[REVISED] ",
            WD_COLOR_INDEX.YELLOW,
            red,
        )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
